#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @author: Rebort

"""
文档文本提取工具支持PDF、DOCX、TXT、MD等格式
"""
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """文档文本提取器"""
    
    @staticmethod
    def extract_text(file_path: str, file_type: str) -> Optional[str]:
        """
        提取文档文本
        
        Args:
            file_path: 文件路径
            file_type: 文件类型 (pdf/docx/txt/md)
            
        Returns:
            提取的文本内容，失败返回None
        """
        try:
            if file_type == 'txt' or file_type == 'md':
                return DocumentExtractor._extract_text_file(file_path)
            elif file_type == 'pdf':
                return DocumentExtractor._extract_pdf(file_path)
            elif file_type == 'docx':
                return DocumentExtractor._extract_docx(file_path)
            else:
                logger.warning(f"不支持的文件类型: {file_type}")
                return None
        except Exception as e:
            logger.error(f"提取文档文本失败: {e}")
            return None
    
    @staticmethod
    def _extract_text_file(file_path: str) -> str:
        """提取TXT/MD文件文本"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"读取文本文件失败 ({encoding}): {e}")
                continue
        
        # 如果所有编码都失败，尝试二进制读取并忽略错误
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"读取文本文件失败: {e}")
            return ""
    
    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """提取PDF文件文本"""
        try:
            import PyPDF2
            
            text_content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            return '\n\n'.join(text_content)
        except ImportError:
            logger.warning("PyPDF2未安装，无法提取PDF文本。请安装: pip install PyPDF2")
            return "PDF文本提取需要安装PyPDF2库"
        except Exception as e:
            logger.error(f"提取PDF文本失败: {e}")
            return f"PDF文本提取失败: {str(e)}"
    
    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """提取DOCX文件文本"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n\n'.join(text_content)
        except ImportError:
            logger.warning("python-docx未安装，无法提取DOCX文本。请安装: pip install python-docx")
            return "DOCX文本提取需要安装python-docx库"
        except Exception as e:
            logger.error(f"提取DOCX文本失败: {e}")
            return f"DOCX文本提取失败: {str(e)}"
